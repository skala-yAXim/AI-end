import json
import re
import sys
import os
from typing import Optional, Dict, List
import pandas as pd
from docx import Document
from langchain_openai import OpenAIEmbeddings
from langchain_qdrant import QdrantVectorStore
from qdrant_client import QdrantClient
from qdrant_client.http.models import Distance, VectorParams

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
from config import OPENAI_API_KEY, DATA_DIR

client = QdrantClient(
    host="localhost",
    port=6333
)

embedding_model = OpenAIEmbeddings(model="text-embedding-3-small")

collection_name = "docs-collection"

# 기존 컬렉션 삭제 및 재생성
if client.collection_exists(collection_name):
    print(f"기존 컬렉션 '{collection_name}' 삭제 중...")
    client.delete_collection(collection_name)
    
client.create_collection(
    collection_name=collection_name,
    vectors_config=VectorParams(size=1536, distance=Distance.COSINE),
)
print(f"새 컬렉션 '{collection_name}' 생성 완료")

qdrant_store = QdrantVectorStore(
    client=client,
    collection_name=collection_name,
    embedding=embedding_model
)

def extract_file_content(filename: str, docs_folder_path: str = None) -> Optional[str]:
    """실제 파일에서 텍스트 콘텐츠를 추출합니다."""
    if docs_folder_path is None:
        docs_folder_path = os.path.join(DATA_DIR, "docs_files")
    
    file_path = os.path.join(docs_folder_path, filename)
    
    if not os.path.exists(file_path):
        return None
        
    file_ext = filename.lower().split('.')[-1]
    
    try:
        if file_ext == 'docx':
            return _extract_from_docx(file_path)
        elif file_ext == 'xlsx':
            return _extract_from_xlsx(file_path)
        elif file_ext == 'txt':
            return _extract_from_txt(file_path)
        else:
            return f"지원하지 않는 파일 형식: {file_ext}"
    except Exception as e:
        return f"파일 읽기 오류: {str(e)}"

def _extract_from_docx(file_path: str) -> str:
    """DOCX 파일에서 텍스트 추출"""
    doc = Document(file_path)
    text_parts = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())
            
    # 표 내용도 추출
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                text_parts.append(" | ".join(row_text))
    
    return "\n".join(text_parts)

def _extract_from_xlsx(file_path: str) -> str:
    """XLSX 파일에서 텍스트 추출"""
    text_parts = []
    
    # 모든 시트 읽기
    excel_file = pd.ExcelFile(file_path)
    for sheet_name in excel_file.sheet_names:
        df = pd.read_excel(file_path, sheet_name=sheet_name)
        
        text_parts.append(f"[시트: {sheet_name}]")
        
        # 컬럼명 추가
        if not df.empty:
            text_parts.append("컬럼: " + " | ".join(str(col) for col in df.columns))
            
            # 데이터 행 추가 (최대 100행)
            for idx, row in df.head(100).iterrows():
                row_text = " | ".join(str(val) for val in row.values if pd.notna(val))
                if row_text.strip():
                    text_parts.append(row_text)
    
    return "\n".join(text_parts)

def _extract_from_txt(file_path: str) -> str:
    """TXT 파일에서 텍스트 추출"""
    encodings = ['utf-8', 'cp949', 'euc-kr']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    
    # 모든 인코딩 실패 시
    return "텍스트 파일 인코딩 오류"

def extract_texts_from_json(json_data):
    docs = []  # VectorDB에 저장할 문서
    for doc in json_data["docs_files"]:
        parsed_doc = parse_json(doc)
        if parsed_doc:  # None이 아닌 경우만 추가
            docs.append(parsed_doc)

    # VectorDB에 저장 (LangChain 방식 복귀)
    if docs:
        qdrant_store.add_texts(
            texts=[doc["text"] for doc in docs],
            metadatas=[doc["metadata"] for doc in docs]
        )
        
        print(f"{len(docs)}개 문서와 메타데이터가 성공적으로 저장되었습니다.")

    return docs

def parse_json(data: dict) -> dict:
    """
    JSON 데이터를 파싱하여 텍스트와 메타데이터를 추출합니다.
    :param data: 문서 메타데이터 JSON
    :return: 파싱된 텍스트와 메타데이터를 포함하는 딕셔너리
    """
    
    filename = data.get("filename", "")
    authors = data.get("author", [])
    last_modified = data.get("last_modified", "")
    file_type = data.get("type", "")
    file_size = data.get("size", 0)
    
    # 실제 파일 콘텐츠 추출 시도
    file_content = extract_file_content(filename)
    
    text_parts = []
    
    # 실제 파일 콘텐츠만 추가 (메타데이터 제거)
    if file_content:
        text_parts.append(file_content)
        has_content = True
    else:
        text_parts.append(f"[파일 내용을 읽을 수 없음: {filename}]")
        has_content = False
    
    combined_text = "\n".join(text_parts)
    
    return {
        "text": combined_text.strip(),
        "metadata": {
            "filename": filename,
            "author": authors,  # 원본 key 유지
            "last_modified": last_modified,  # 원본 key 유지
            "type": file_type,  # 원본 key 유지
            "size": file_size,  # 원본 key 유지
            "url": data.get("url")  # 원본 key 유지
        }
    }

# 사용 예시
with open(DATA_DIR + "/docs_files_data.json", encoding="utf-8") as f:  # 파일명 변경
    json_data = json.load(f)

parsed_docs = extract_texts_from_json(json_data)

# 출력 예시
for d in parsed_docs:
    print("-----")
    print("텍스트 내용:")
    print(d["text"][:200] + "..." if len(d["text"]) > 200 else d["text"])  # 처음 200자만 출력
    print("메타데이터:")
    print(d["metadata"])
